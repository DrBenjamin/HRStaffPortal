##### `/pages/functions/functions.py`
##### HR Staff Portal
##### Open-Source, hosted on https://github.com/DrBenjamin/HRStaffPortal
##### Please reach out to benjamin.gross@giz.de for any questions
#### Loading needed Python libraries
import streamlit as st
import platform
import pandas as pd
import io
import os
import xlsxwriter
import openpyxl
import xlrd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from streamlit_qrcode_scanner import qrcode_scanner
import qrcode
import sys
sys.path.insert(1, "pages/functions/")
from network import downzip

    
    
    
#### All shared general functions
### Function: check_password = Password / user checking
def check_password():
    # Session states
    if ("username" not in st.session_state):
        st.session_state["username"] = ''
    if ("password" not in st.session_state):
        st.session_state["password"] = ''
    if ("admin" not in st.session_state):
        st.session_state["admin"] = False
    if ("password_correct" not in st.session_state):
        st.session_state["password_correct"] = False
    if ('logout' not in st.session_state):
        st.session_state['logout'] = False

    # Checks whether a password entered by the user is correct
    def password_entered():
        try:
            if st.session_state["username"] in st.secrets["passwords"] and st.session_state["password"] == st.secrets["passwords"][st.session_state["username"]]:
                st.session_state["password_correct"] = True
                st.session_state["admin"] = False
    
                # Delete username + password
                del st.session_state["password"]
                del st.session_state["username"]

            # Checks whether a password entered by the user is correct for admins
            elif st.session_state["username"] in st.secrets["admins"] and st.session_state["password"] == st.secrets["admins"][st.session_state["username"]]:
                st.session_state["password_correct"] = True
                st.session_state["admin"] = True
    
                # Delete username + password
                del st.session_state["password"]
                del st.session_state["username"]
    
            # No combination fits
            else:
                st.session_state["password_correct"] = False
        except Exception as e:
            print('Exception in `password_entered` function. Error: ', e)
            st.session_state["password_correct"] = False


    ## Sidebar
    # Download images
    if not os.path.exists(st.secrets['custom']['images_path'] + st.secrets['custom']['sidebar_image']):
        downzip(st.secrets['custom']['images_url'], [st.secrets['custom']['images_zip']], st.secrets['custom']['images_path'])


    # Show Sidebar Header Image
    st.sidebar.image(st.secrets['custom']['sidebar_image'])

    # Header switch
    if st.session_state['header'] == True:
        index = 0
    elif st.session_state['header'] == False:
        index = 1
    else:
        index = 0
    header = st.sidebar.radio(label = 'Switch headers on or off', options = ('on', 'off'), index = index,
                              horizontal = True)
    if header == 'on':
        st.session_state['header'] = True
    elif header == 'off':
        st.session_state['header'] = False
    else:
        st.session_state['header'] = True

    # First run, show inputs for username + password
    if "password_correct" not in st.session_state:
        st.sidebar.subheader('Please enter username and password')
        st.sidebar.text_input(label = "Username", on_change = password_entered, key = "username")
        st.sidebar.text_input(label = "Password", type = "password", on_change = password_entered, key = "password")
        return False

    # Password not correct, show input + error
    elif not st.session_state["password_correct"]:
        st.sidebar.text_input(label = "Username", on_change = password_entered, key = "username")
        st.sidebar.text_input(label = "Password", type = "password", on_change = password_entered, key = "password")
        if (st.session_state['logout']):
            st.sidebar.success('Logout successful!', icon = "✅")
        else:
            st.sidebar.error(body = "User not known or password incorrect!", icon = "🚨")
        return False

    else:
        # Password correct
        st.sidebar.success(body = 'You are logged in.', icon = "✅")
        st.sidebar.info(body = 'You can close this menu')
        st.sidebar.button(label = 'Logout', on_click = logout)
        return True



### Function logger = Logging debug messages to file
def logger():
    logname = "files/file_1.log"
    logger.add(logname)
    logger.debug("This is a debug message")



### Funtion: logout = Logout button
def logout():
    # Set `logout` to get logout-message
    st.session_state['logout'] = True

    # Set password to `false`
    st.session_state["password_correct"] = False



### Function: import_excel = MS Excel File (xlsx) to Pandas dataframe
def import_excel():
    data = None
    output = [[]]
    output2 = [[]]
    uploaded_file = st.file_uploader("Choose an Excel document for data import", type = ['xls', 'xlsx', 'doc', 'docx'])
    if uploaded_file is not None:
        # To read file as bytes:
        bytes_data = uploaded_file.getvalue()
        suffix = uploaded_file.name.split('.')[-1]
        
        # Write Excel to dataframe
        try:
            if suffix == 'xls' or suffix == 'xlsx':
                data = pd.read_excel(io.BytesIO(bytes_data), sheet_name = 0)
                for index, row in data.iterrows():
                    if str(row[2]) != 'nan' and str(row[2]) != '':
                        output.append([index, str(row[1]).strip()])
                output = pd.DataFrame(output, columns = ['ID', 'Position'])
                output = output.set_index('ID')
                output = output.drop_duplicates(ignore_index = True)
                output = output.drop(0, axis = 0)
                output2 = pd.DataFrame('Nothing')
            elif suffix == 'doc' or suffix == 'docx':
                with open("temp." + suffix, "wb") as f:
                    f.write(bytes_data)
                document = Document('temp.' + suffix)
                docx = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
                data = pd.DataFrame(docx)
                out1 = []
                out2 = []
                empty_bool = False
                for index, row in data.iterrows():
                    if index == 0:
                        out1.append(row[1].strip().title())
                    else:
                        if row[1] == '':
                            empty_bool = True
                            continue
                        else:
                            if empty_bool == True:
                                empty_bool = False
                                out1.append(row[1].strip().title())
                            else:
                                out2.append(row[1].strip().title())
                output = pd.DataFrame(out1, columns = ['Department'])
                output = output.drop_duplicates(ignore_index = True)
                output.index = output.index + 1
                output2 = pd.DataFrame(out2, columns = ['Unit'])
                output2 = output2.drop_duplicates(ignore_index = True)
                output2.index = output2.index + 1
        except:
            print('No import data present')
    return output, output2
    
    
    
### Function: export_excel = Pandas dataframe to MS Excel Makro File (xlsm)
def export_excel(sheet, column, columns, length, data,
                 sheet2 = 'N0thing', column2 = 'A', columns2 = '', length2 = '', data2 = '',
                 sheet3 = 'N0thing', column3 = 'A', columns3 = '', length3 = '', data3 = '',
                 sheet4 = 'N0thing', column4 = 'A', columns4 = '', length4 = '', data4 = '',
                 sheet5 = 'N0thing', column5 = 'A', columns5 = '', length5 = '', data5 = '',
                 sheet6 = 'N0thing', column6 = 'A', columns6 = '', length6 = '', data6 = '',
                 sheet7 = 'N0thing', column7 = 'A', columns7 = '', length7 = '', data7 = '',
                 image = 'NoImage', image_pos = 'D1', excel_file_name = 'Export.xlsm'):


    ## Store fuction arguments in array
    # Create empty array
    func_arr = []

    # Add function arguments to array
    func_arr.append([sheet, column, columns, length, data])
    func_arr.append([sheet2, column2, columns2, length2, data2])
    func_arr.append([sheet3, column3, columns3, length3, data3])
    func_arr.append([sheet4, column4, columns4, length4, data4])
    func_arr.append([sheet5, column5, columns5, length5, data5])
    func_arr.append([sheet6, column6, columns6, length6, data6])
    func_arr.append([sheet7, column7, columns7, length7, data7])


    ## Create an Excel file filled with a pandas dataframe using XlsxWriter as engine
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine = 'xlsxwriter') as writer:
        for i in range(7):
            if (func_arr[i][0] != 'N0thing'):
                # Add dataframe data to worksheet
                func_arr[i][4].to_excel(writer, sheet_name = func_arr[i][0], index = False)

                # Define worksheet
                worksheet = writer.sheets[func_arr[i][0]]

                # Add a table to the worksheet
                if func_arr[i][1] != 'A':
                    span = "A1:%s%s" % (func_arr[i][1], func_arr[i][3])
                    worksheet.add_table(span, {'columns': func_arr[i][2]})
                    range_table = "A:" + func_arr[i][1]
                    worksheet.set_column(range_table, 30)

                # Add image to worksheet
                if (image != 'NoImage'):
                    # Saving image as png to a buffer
                    # byteIO = io.BytesIO()
                    # image.save(byteIO, format = 'PNG')
                    # pic = byteIO.getvalue()

                    # Saving image as png temp file
                    f = open('files/temp.png', 'wb')
                    f.write(image)
                    f.close()

                    # Insert in worksheet
                    worksheet.insert_image(image_pos, 'files/temp.png')


        ## Add Excel VBA code
        workbook = writer.book
        workbook.add_vba_project('files/vbaProject.bin')


        ## Saving changes
        workbook.close()
        writer.save()


        ## Download Button
        st.download_button(label = 'Download Excel document', data = buffer, file_name = excel_file_name,
                           mime = "application/vnd.ms-excel.sheet.macroEnabled.12")



### Function: export_docx = Pandas dataframe to MS Word file (docx)
def export_docx(data, faq, docx_file_name = 'Handbook.docx'):
    document = Document()


    ## Sorting dataframe by chapter and paragraph
    data = data.sort_values(['HANDBOOK_CHAPTER', 'HANDBOOK_PARAGRAPH'], ascending = [True, True])


    ## Writing handbook
    # Adding handbook header
    document.add_heading('User Handbook', 0)

    # Adding table of contents
    document.add_heading('Table of contents', level = 1)
    paragraph = document.add_paragraph()
    chapter = 0
    paragraf = 0
    for i in range(len(data)):
        if (data.iloc[i]['HANDBOOK_PARAGRAPH'] != paragraf):
            if (chapter != data.iloc[i]['HANDBOOK_CHAPTER']):
                paragraph.add_run('\n' + data.iloc[i]['HANDBOOK_CHAPTER_DESCRIPTION'] + '\n').bold = True
                chapter += 1
            if (str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1] == '0'):
                paragraph.add_run(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[0] + ' - ' + data.iloc[i][
                    'HANDBOOK_TEXT_HEADLINE'] + '\n')
            else:
                if len(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1]) == 1:
                    placer = '\t'
                elif len(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1]) == 2:
                    placer = '\t\t'
                elif len(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1]) == 3:
                    placer = '\t\t\t'
                else:
                    placer = '\t\t\t\t'
                paragraph.add_run(placer + str(data.iloc[i]['HANDBOOK_PARAGRAPH']) + ' - ' + data.iloc[i][
                    'HANDBOOK_TEXT_HEADLINE'] + '\n')
        paragraf = data.iloc[i]['HANDBOOK_PARAGRAPH']

    # Writing paragraphs
    chapter = 0
    paragraf = 0
    for i in range(len(data)):
        # Adding chapter headings
        if data.iloc[i]['HANDBOOK_CHAPTER'] > chapter:
            document.add_page_break()

            # Adding handbook header
            document.add_heading('User Handbook', 0)
            document.add_heading(data.iloc[i]['HANDBOOK_CHAPTER_DESCRIPTION'], level = 1)
            paragraph = document.add_paragraph()
            paragraph.add_run(data.iloc[i]['HANDBOOK_CHAPTER_TEXT']).italic = True
            chapter = data.iloc[i]['HANDBOOK_CHAPTER']

        # Adding paragraph headings
        if (data.iloc[i]['HANDBOOK_PARAGRAPH'] != paragraf):
            if len(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1]) == 1:
                if str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1] == '0':
                    document.add_heading(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[0] + '\t' + data.iloc[i][
                        'HANDBOOK_TEXT_HEADLINE'], level = 2)
                else:
                    document.add_heading(
                        str(data.iloc[i]['HANDBOOK_PARAGRAPH']) + '\t' + data.iloc[i]['HANDBOOK_TEXT_HEADLINE'],
                        level = 2)
            elif len(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1]) == 2:
                document.add_heading(
                    str(data.iloc[i]['HANDBOOK_PARAGRAPH']) + '\t' + data.iloc[i]['HANDBOOK_TEXT_HEADLINE'], level = 3)
            elif len(str(data.iloc[i]['HANDBOOK_PARAGRAPH']).split('.')[1]) == 3:
                document.add_heading(
                    str(data.iloc[i]['HANDBOOK_PARAGRAPH']) + '\t' + data.iloc[i]['HANDBOOK_TEXT_HEADLINE'], level = 4)
            else:
                document.add_heading(
                    str(data.iloc[i]['HANDBOOK_PARAGRAPH']) + '\t' + data.iloc[i]['HANDBOOK_TEXT_HEADLINE'], level = 5)

        # Adding paragraphs
        paragraph = document.add_paragraph()
        if (data.iloc[i]['HANDBOOK_PARAGRAPH'] != paragraf):
            if (data.iloc[i]['HANDBOOK_PARAGRAPH_TEXT'] != None):
                paragraph.add_run(data.iloc[i]['HANDBOOK_PARAGRAPH_TEXT'] + '\n\n').italic = True
        paragraph.add_run(data.iloc[i]['HANDBOOK_TEXT'])
        paragraf = data.iloc[i]['HANDBOOK_PARAGRAPH']

        # Adding image
        if (data.iloc[i]['HANDBOOK_IMAGE_TEXT'] != 'Placeholder image.'):
            save_img(data = data.iloc[i]['HANDBOOK_IMAGE'], filename = 'images/temp.png')
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            run.add_picture('images/temp.png')
            paragraph.add_run('\n' + data.iloc[i]['HANDBOOK_IMAGE_TEXT']).italic = True


    ## Writing FAQ
    # Adding handbook header
    document.add_page_break()
    document.add_heading('User Handbook', 0)

    # Adding FAQ header
    document.add_heading('FAQ', level = 1)

    # Adding FAQ items
    for i in range(len(faq)):
        paragraph = document.add_paragraph()
        paragraph.add_run('Question: ').bold = True
        paragraph.add_run(faq[i][0].upper() + '\n')
        paragraph.add_run('Category & Sub-Category: ').bold = True
        paragraph.add_run(faq[i][2] + ' / ' + faq[i][3] + '\n')
        paragraph.add_run('Ben`s answer: ').bold = True
        paragraph.add_run(faq[i][1] + '\n\n')


    ## Create a Word file using python-docx as engine
    buffer = io.BytesIO()
    document.save(buffer)
    if os.path.exists("images/temp.png"):
        os.remove("images/temp.png")


    ## Download Button
    st.download_button(label = 'Download Word document', data = buffer, file_name = docx_file_name,
                       mime = "application/vnd.openxmlformats")



### Function: loadFile = Converts digital data to binary format
def load_file(filename):
    with open(filename, 'rb') as file:
        binaryData = file.read()
    return binaryData



### Function: save_img = Converts binary image data to png file
def save_img(data, filename = 'temp.png'):
    file = open(filename, 'wb')
    file.write(data)
    file.close()



### Function: generateID = Generates a 5-digits ID
def generateID(id):
    if (id < 10):
        generated_id = '0000' + str(id)
    elif (id < 100):
        generated_id = '000' + str(id)
    elif (id < 1000):
        generated_id = '00' + str(id)
    elif (id < 10000):
        generated_id = '0' + str(id)

    # Return the 5-digit ID
    return (generated_id)



### Function: generate_qrcode = QR Code generator
def generate_qrcode(data):
    # Encoding data using make() function
    image = qrcode.make(data)

    # Saving image as png in a buffer
    byteIO = io.BytesIO()
    image.save(byteIO, format = 'PNG')

    # Return qrcode
    return byteIO.getvalue()



### parse_national_id = Parsing National ID QR Code data to list
def parse_national_id(text):
    val = text.split('~')
    if len(val) == 12:
        if len(val[6].split(', ')) > 1:
            fname, mname = val[6].split(', ')
        else:
            fname = val[6]
            mname = ''
        lname = val[4]
        gender = str(val[8]).upper()
        raw_dob = val[9]
        nat_id = str(val[5])

        # Cleaning data format
        dateOfBirth = str(raw_dob).split(" ")
        day = dateOfBirth[0]
        year = dateOfBirth[2]
        month = dateOfBirth[1]
        month_var = {"JAN": 1, "FEB": 2, "MAR": 3, "APR": 4, "MAY": 5, "JUN": 6, "JUL": 7, "AUG": 8, "SEP": 9,
                     "OCT": 10, "NOV": 11, "DEC": 12}
        num_month = month_var[month.upper()]
        dob = year + "-" + month + "-" + day
        result = {"first_name": fname, "middle_name": mname, "last_name": lname, "gender": gender, "nation_id": nat_id,
                  "dob": dob}

        # Return result
        return result



### Function: qrcode_reader = Scans National IDs QR Code
def qrcode_reader():
    qrcode = None
    qrcode = qrcode_scanner(key = 'qrcode_scanner')
    return qrcode



### Function build_employee = Creates list of names of confirmed / not confirmed
def build_employees(data, not_confirmed, confirmed):
    row = []
    names = []
    not_in_list = True
    already_in = ''
    if not_confirmed is not None:
        already_in = not_confirmed.split(' ')
    elif confirmed is not None:
        already_in += confirmed.split(' ')
    already_in = list(filter(None, already_in))
    for row in data:
        for eno in already_in:
            if row[3] == eno:
                not_in_list = False
        if not_in_list == True:
            if row[5] is None:
                names.append(str(row[1] + ' ' + row[2] + ', ' + row[3] + ', ' + row[4] + ' ('')'))
            else:
                names.append(str(row[1] + ' ' + row[2] + ', ' + row[3] + ', ' + row[4] + ' (' + row[5] + ')'))
        else:
            not_in_list = True
    return names, already_in



### Function rebuild_confirmation = Creates new lists of of confirmed / not confirmed
def rebuild_confirmation(option_attendee, confirmed, not_confirmed, employee):
    not_confirmed_after = ''
    not_confirmed_before = not_confirmed
    if option_attendee > 0:
        not_confirmed_after = not_confirmed_before.replace(employee, '')
        if confirmed != None:
            confirmed += ' ' + employee
        else:
            confirmed = employee
    else:
        if not_confirmed != None:
            not_confirmed_after = not_confirmed
    return confirmed, not_confirmed_after



### Function confirmed_query = Creates query to get just confirmed from database
def confirmed_query(confirmed):
    confirmed = confirmed.split(' ')
    confirmed_query = ""
    for row in confirmed:
        if len(confirmed_query) == 0:
            confirmed_query = confirmed_query + "EMPLOYEE_NO = " + "'" + row + "'"
        else:
            if row != '':
                confirmed_query = confirmed_query + " OR EMPLOYEE_NO = " + "'" + row + "'"
    query = "SELECT FORENAME, SURNAME, EMPLOYEE_NO, IMAGE FROM `idcard`.`IMAGEBASE` WHERE %s;" % (confirmed_query)
    return query



### Function header = Shows header information 
def header(title, data_desc, expanded = True):
    with st.expander("Header", expanded = expanded):
        ## Header information
        st.title(title)
        st.image(st.secrets['custom']['facility_image'])
        st.header(st.secrets['custom']['facility'] + ' (' + st.secrets['custom']['facility_abbreviation'] + ')')
        st.subheader(st.secrets['custom']['facility_abbreviation'] + ' ' + data_desc)
        st.write('All data is stored in a local MySQL databank on a dedicated Server hosted at ' + st.secrets['custom'][
            'facility_abbreviation'] + '.')
        st.write(
            'The ' + title + ' is developed with Python (v' + platform.python_version() + ') and Streamlit (v' + st.__version__ + ').')



### Function: landing_page = Shows the landing page (not loged in state)
def landing_page(page):
    ## Title and information
    header = 'Welcome to ' + page
    st.title(header)
    st.header(st.secrets['custom']['facility'] + ' (' + st.secrets['custom']['facility_abbreviation'] + ')')
    st.subheader('User Login')
    st.info(body = 'Please login (sidebar on the left) to access the ' + page, icon = "ℹ️")


    ## Sub-pages menu
    st.subheader('Pages without login')
    st.write('You can access these pages without being logged in:')
    st.write("<a href='Handbook' target='_self'>Handbook</a>", unsafe_allow_html = True)
    st.write("<a href='Changelog' target='_self'>Changelog</a>", unsafe_allow_html = True)

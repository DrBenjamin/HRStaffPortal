##### `pages/7 - 📖_Handbook_Export.py`
##### Kamazu Central Hospital (KCH) HR Staff Portal Prototype
##### Open-Source, hostet on https://github.com/DrBenjamin/HRStaffPortal
##### Please reach out to benjamin.gross@giz.de for any questions
#### Loading needed Python libraries
import streamlit as st
import streamlit.components.v1 as stc
import mysql.connector
from docx import Document
from docx.shared import Inches
import sys
sys.path.insert(1, "pages/functions/")
from functions import loadFile




#### Streamlit initial setup
st.set_page_config(
  page_title = "HR Staff Portal",
  page_icon = "images/thumbnail.png",
  layout = "centered",
  initial_sidebar_state = "expanded",
  menu_items = { 
         'Get Help': 'http://www.health.gov.mw/index.php/contact-moh/head-office',
         'Report a bug': "http://www.health.gov.mw/index.php/contact-moh/head-office",
         'About': "This is the HR Staff Portal Version 0.2.0"
        }
)




#### Functions
### Function: run_query = Initial SQL Connection
def init_connection():
  ## Initialize connection
  try:
    return mysql.connector.connect(**st.secrets["mysql_benbox"])
  except:
    print("An exception occurred in function `init_connection`")
    st.error(body = 'Databank connection timeout!', icon = "🚨")
    st.stop()



### Function: run_query = SQL query
def run_query(query):
  with conn.cursor() as cur:
    # Perform query
    try:
      cur.execute(query)
      return cur.fetchall()
    
    except:
      print('An exception occurred in function `run_query` with query \"' + query + '\"')
      
      
      
### Function: pictureUploader = uploads handbook images
def pictureUploader(image, index):
  # Initialize connection
  connection = mysql.connector.connect(**st.secrets["mysql_benbox"])
  cursor = connection.cursor()
  
  # SQL statement
  sql_insert_blob_query = """ UPDATE `HANDBOOK_USER` SET HANDBOOK_IMAGE = %s WHERE ID = %s;"""
  
  # Convert data into tuple format
  insert_blob_tuple = (image, index)
  result = cursor.execute(sql_insert_blob_query, insert_blob_tuple)
  connection.commit()
      
      
      
### Function: lastID = checks for last ID number in Table (to add data after)
def lastID(url):
  query = "SELECT MAX(ID) FROM %s;" %(url)
  rows = run_query(query)
  
  # Check for ID
  for row in rows:
    if (row[0] != None):
      id = int(row[0]) + 1
    else:
      id = 1
      break
  
  # Return ID    
  return id



### Function: generateID = Generates an 5-digits ID
def generateID(id):
  if (id < 10):
    generated_id = '0000' + str(id)
  elif (id < 100):
    generated_id = '000' + str(id)
  elif (id < 1000):
    generated_id = '00' + str(id)
  elif (id < 10000):
    generated_id = '0' + str(id)
    
  # Return the 5-digit ID
  return(generated_id)


  

#### Sidebar
### Handbook Sidebar
## Sidebar Header Image
st.sidebar.image('images/MoH.png')




#### Main Program
### Header
## Header
st.title('Handbook Export')
st.subheader('User Handbook Export')



### Handbook
## Open databank connection
conn = init_connection()


## Docx Export
document = Document()

document.add_heading('User Handbook', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')
